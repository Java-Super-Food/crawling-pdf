import os
import base64
from pdf2image import convert_from_path
from openai import OpenAI
from docx import Document
from io import BytesIO
from dotenv import load_dotenv
import subprocess
import re
from PIL import Image

load_dotenv()
Image.MAX_IMAGE_PIXELS = None


# PDF → IMAGE RENDERER (PER PAGE)
class PDFToImageRenderer:
    def __init__(self, pdf_path: str, dpi: int = 120):
        self.pdf_path = pdf_path
        self.dpi = dpi
        self.total_pages = self._get_total_pages()

    def _get_total_pages(self) -> int:
        result = subprocess.run(
            ["pdfinfo", self.pdf_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )

        match = re.search(r"Pages:\s+(\d+)", result.stdout)
        if not match:
            raise RuntimeError("Gagal membaca jumlah halaman PDF")

        return int(match.group(1))

    def render_pages(self, start_page: int = 1, max_pages: int | None = None):
        if start_page > self.total_pages:
            raise ValueError("start_page melebihi total halaman PDF")

        if max_pages:
            last_page = min(start_page + max_pages - 1, self.total_pages)
        else:
            last_page = self.total_pages

        for page in range(start_page, last_page + 1):
            print(f"📄 Rendering page {page}...")
            images = convert_from_path(
                self.pdf_path,
                dpi=self.dpi,
                first_page=page,
                last_page=page
            )
            yield page, images[0]


# OPENAI VISION OCR
class OpenAIVisionOCR:
    def __init__(self, model: str = "gpt-4.1"):
        self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.model = model

    def _normalize_image(self, image, max_width=1800):
        image = image.convert("L")

        if image.width > max_width:
            ratio = max_width / image.width
            image = image.resize(
                (max_width, int(image.height * ratio))
            )

        return image

    def _image_to_base64(self, image) -> str:
        buffer = BytesIO()
        image.save(buffer, format="PNG", optimize=True)
        return base64.b64encode(buffer.getvalue()).decode()

    def extract_catalogue_data(self, image) -> dict:
        image = self._normalize_image(image)
        image_base64 = self._image_to_base64(image)

        prompt = """
Ini adalah SATU halaman katalog produk.

Tugas kamu:
- Lakukan OCR dari gambar
- Abaikan desain, ornamen, background
- Abaikan watermark atau slogan umum
- Fokus pada INTI PRODUK
- Jika teks tidak lengkap, buat deskripsi profesional yang masuk akal

Output HARUS format PERSIS:
Title: ...
Description: ...
"""

        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": prompt},
                        {
                            "type": "input_image",
                            "image_url": f"data:image/png;base64,{image_base64}"
                        }
                    ]
                }
            ]
        )

        return self._parse_response(response.output_text)

    def _parse_response(self, text: str) -> dict:
        title = ""
        description_lines = []

        for line in text.splitlines():
            if line.lower().startswith("title"):
                title = line.replace("Title:", "").strip()
            elif line.lower().startswith("description"):
                description_lines.append(
                    line.replace("Description:", "").strip()
                )
            else:
                description_lines.append(line.strip())

        return {
            "title": title,
            "description": " ".join(description_lines)
        }


# WORD DOCUMENT BUILDER (APPEND MODE)
class WordCatalogueBuilder:
    def __init__(self, output_path: str):
        self.output_path = output_path

        if os.path.exists(output_path):
            self.doc = Document(output_path)
            print("📄 Existing DOCX loaded (append mode)")
        else:
            self.doc = Document()
            print("📄 New DOCX created")

    def add_product(self, page_number: int, data: dict):
        if not data["title"]:
            return

        self.doc.add_heading(
            f"Page {page_number} – {data['title']}",
            level=2
        )
        self.doc.add_paragraph(data["description"])

    def save(self):
        self.doc.save(self.output_path)


# MAIN APPLICATION (BATCH SAFE)
class CatalogueOCRApp:
    def __init__(
        self,
        pdf_path: str,
        output_docx: str,
        start_page: int = 1,
        max_pages: int | None = None
    ):
        self.renderer = PDFToImageRenderer(pdf_path)
        self.ocr = OpenAIVisionOCR()
        self.writer = WordCatalogueBuilder(output_docx)
        self.start_page = start_page
        self.max_pages = max_pages

    def run(self):
        for page_number, image in self.renderer.render_pages(
            start_page=self.start_page,
            max_pages=self.max_pages
        ):
            print(f"🔍 OCR processing page {page_number}...")

            try:
                data = self.ocr.extract_catalogue_data(image)
                self.writer.add_product(page_number, data)
                self.writer.save()
                print(f"✅ Page {page_number} saved")
            except Exception as e:
                print(f"❌ Error on page {page_number}: {e}")
                continue

        print("🏁 OCR batch finished safely")


# ENTRY POINT
if __name__ == "__main__":
    app = CatalogueOCRApp(
        pdf_path="catalog.pdf",
        output_docx="catalogue_output.docx",
        start_page=6,
        max_pages=50
    )
    app.run()
